#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
基于学校模板的 Markdown -> DOCX 转换器

目标：
1. 基于模板填充（保留封面/声明页等前置内容）
2. 自动处理标题、正文、目录、表格、图片
3. 生成可提交的论文文档（仍建议最终人工核对目录页码）

默认输入输出：
- 输入：06-毕业论文.md
- 模板：2西南科技大学本科毕业论文（设计）模板2026版.docx
- 输出：毕业论文_最终版.docx
"""

from __future__ import annotations

import re
import subprocess
from pathlib import Path
from typing import List, Optional

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


# -----------------------------
# 基础工具
# -----------------------------


def set_run_font(run, cn: str = "宋体", en: str = "Times New Roman", size_pt: float = 12, bold: Optional[bool] = None):
    run.font.name = en
    run._element.rPr.rFonts.set(qn("w:eastAsia"), cn)
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold


def set_exact_22(paragraph):
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(22)


def style_heading_chapter(paragraph):
    set_outline_level(paragraph, 1)
    set_exact_22(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.space_before = Pt(7.8)
    paragraph.paragraph_format.space_after = Pt(7.8)
    for run in paragraph.runs:
        set_run_font(run, cn="黑体", en="Times New Roman", size_pt=18, bold=True)


def style_heading_section(paragraph):
    set_outline_level(paragraph, 2)
    set_exact_22(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.space_before = Pt(7.8)
    paragraph.paragraph_format.space_after = Pt(7.8)
    for run in paragraph.runs:
        set_run_font(run, cn="黑体", en="Times New Roman", size_pt=14, bold=True)


def style_heading_subsection(paragraph):
    set_outline_level(paragraph, 3)
    set_exact_22(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.space_before = Pt(7.8)
    paragraph.paragraph_format.space_after = Pt(7.8)
    for run in paragraph.runs:
        set_run_font(run, cn="黑体", en="Times New Roman", size_pt=12, bold=True)


def style_body(paragraph, indent: bool = True):
    set_exact_22(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.first_line_indent = Cm(0.74) if indent else None
    for run in paragraph.runs:
        set_run_font(run, cn="宋体", en="Times New Roman", size_pt=12, bold=False)


def set_outline_level(paragraph, level: int):
    """
    设置段落大纲级别（1-9），用于 Word 自动目录识别（TOC \\u）。
    """
    if level < 1 or level > 9:
        return
    ppr = paragraph._p.get_or_add_pPr()
    outline = ppr.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        ppr.append(outline)
    outline.set(qn("w:val"), str(level - 1))


def style_caption(paragraph, is_figure: bool):
    set_exact_22(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    for run in paragraph.runs:
        set_run_font(run, cn="宋体", en="Times New Roman", size_pt=10.5, bold=is_figure)


def style_title_cn(paragraph):
    set_exact_22(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.space_before = Pt(7.8)
    paragraph.paragraph_format.space_after = Pt(7.8)
    for run in paragraph.runs:
        set_run_font(run, cn="黑体", en="Times New Roman", size_pt=24, bold=False)


def style_title_en(paragraph):
    set_exact_22(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.space_before = Pt(7.8)
    paragraph.paragraph_format.space_after = Pt(7.8)
    for run in paragraph.runs:
        set_run_font(run, cn="黑体", en="Times New Roman", size_pt=24, bold=True)


def style_abstract_heading(paragraph, text: str):
    set_exact_22(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.space_before = Pt(7.8)
    paragraph.paragraph_format.space_after = Pt(7.8)
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(text)
    set_run_font(run, cn="黑体", en="Times New Roman", size_pt=18, bold=True)


def style_keywords(paragraph):
    set_exact_22(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.space_before = Pt(7.8)
    paragraph.paragraph_format.space_after = Pt(7.8)
    for run in paragraph.runs:
        set_run_font(run, cn="宋体", en="Times New Roman", size_pt=12, bold=False)


def parse_markdown_table_row(row: str) -> List[str]:
    row = row.strip()
    if row.startswith("|"):
        row = row[1:]
    if row.endswith("|"):
        row = row[:-1]
    return [cell.strip() for cell in row.split("|")]


def is_table_separator(line: str) -> bool:
    line = line.strip()
    if not line.startswith("|"):
        return False
    # 例如: |---|---:|:---|
    return bool(re.match(r"^\|[\s:\-\|]+\|$", line))


def remove_body_from_anchor(doc: Document, anchor_text: str):
    """
    从包含 anchor_text 的段落开始，删除其后的正文元素（保留封面/声明页等模板前置内容）。
    """
    anchor_para = None
    for p in doc.paragraphs:
        if anchor_text in p.text:
            anchor_para = p
            break
    if anchor_para is None:
        raise RuntimeError(f"未在模板中找到锚点段落: {anchor_text!r}")

    body = doc._body._element
    children = list(body)
    start_idx = children.index(anchor_para._element)
    for elem in children[start_idx:]:
        # 不删除 section 属性节点
        if elem.tag.endswith("sectPr"):
            continue
        body.remove(elem)


def try_convert_svg_to_png(svg_path: Path) -> Optional[Path]:
    """
    优先使用 rsvg-convert，其次使用 macOS qlmanage 将 SVG 转为 PNG。
    """
    if not svg_path.exists():
        return None
    out_dir = Path("/tmp")
    out_png = out_dir / f"{svg_path.stem}.png"

    # 1) 首选 rsvg-convert（质量稳定、命名可控）
    try:
        subprocess.run(
            ["rsvg-convert", "-w", "1800", "-f", "png", "-o", str(out_png), str(svg_path)],
            check=True,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
        )
        if out_png.exists():
            return out_png
    except Exception:
        pass

    # 2) 兜底 qlmanage（macOS 原生）
    try:
        subprocess.run(
            ["qlmanage", "-t", "-s", "2000", "-o", str(out_dir), str(svg_path)],
            check=True,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
        )
        # qlmanage 输出文件名通常为 <stem>.png
        quicklook_png = out_dir / f"{svg_path.stem}.png"
        if quicklook_png.exists():
            return quicklook_png
    except Exception:
        pass

    return None


def resolve_image_path(raw_path: str, base_dir: Path) -> Optional[Path]:
    p = Path(raw_path)
    if not p.is_absolute():
        # 兼容 markdown 中 `aGraduate/xxx` 路径
        if str(p).startswith("aGraduate/"):
            p = base_dir.parent / p
        else:
            p = base_dir / p
    if p.exists():
        if p.suffix.lower() == ".svg":
            return try_convert_svg_to_png(p)
        return p
    return None


def enable_update_fields_on_open(doc: Document):
    """
    让 Word 打开文档时提示/执行域更新（目录、页码等）。
    """
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def add_toc_field(doc: Document):
    """
    插入 Word 自动目录域（需要在 Word 中更新域后显示最终页码）。
    """
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = None

    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    run = p.add_run()
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
    run._r.append(instr_text)

    run = p.add_run()
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_separate)

    hint = p.add_run("（在 Word 中右键目录并更新域以生成页码）")
    set_run_font(hint, cn="宋体", en="Times New Roman", size_pt=10.5, bold=False)

    run = p.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)

    style_body(p, indent=False)


def extract_sections(lines: List[str]):
    """
    从当前论文 markdown 结构中提取：
    - 题目中英文
    - 中英文摘要与关键词
    - 目录行
    - 正文（从 '# 1 ...' 开始）
    """

    def section_content(h2_title: str) -> List[str]:
        start = None
        for i, line in enumerate(lines):
            if line.strip() == f"## {h2_title}":
                start = i + 1
                break
        if start is None:
            return []
        end = len(lines)
        for j in range(start, len(lines)):
            if lines[j].startswith("## "):
                end = j
                break
        return [x.rstrip("\n") for x in lines[start:end]]

    # 题目区
    title_block = section_content("题目名称")
    title_cn = ""
    title_en = ""
    for i, line in enumerate(title_block):
        t = line.strip()
        if not t:
            continue
        if t == "英文题目：":
            # 下一个非空即英文题目
            for nxt in title_block[i + 1 :]:
                if nxt.strip():
                    title_en = nxt.strip()
                    break
            break
        if not title_cn:
            title_cn = t

    # 中文摘要
    zh_block = section_content("中文摘要")
    zh_keywords = ""
    zh_abs_lines = []
    for line in zh_block:
        t = line.strip()
        if not t or t == "---":
            continue
        if t.startswith("关键词："):
            zh_keywords = t
            continue
        zh_abs_lines.append(t)

    # 英文摘要
    en_block = section_content("Abstract")
    en_keywords = ""
    en_abs_lines = []
    for line in en_block:
        t = line.strip()
        if not t or t == "---":
            continue
        if t.lower().startswith("key words:"):
            en_keywords = t
            continue
        en_abs_lines.append(t)

    # 目录
    toc_block = section_content("目  录")
    toc_lines = [x.strip() for x in toc_block if x.strip() and x.strip() != "---"]

    # 正文：从 '# 1 ' 开始
    body_start = None
    for i, line in enumerate(lines):
        if re.match(r"^#\s+\d+\s+", line.strip()):
            body_start = i
            break
    body_lines = [x.rstrip("\n") for x in lines[body_start:]] if body_start is not None else []

    return {
        "title_cn": title_cn,
        "title_en": title_en,
        "zh_abs_lines": zh_abs_lines,
        "zh_keywords": zh_keywords,
        "en_abs_lines": en_abs_lines,
        "en_keywords": en_keywords,
        "toc_lines": toc_lines,
        "body_lines": body_lines,
    }


def add_table(doc: Document, rows: List[List[str]]):
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    # 有些环境可能没有该 style，不强依赖
    try:
        table.style = "Table Grid"
    except Exception:
        pass

    for r_i, row in enumerate(rows):
        for c_i in range(col_count):
            text = row[c_i] if c_i < len(row) else ""
            cell = table.cell(r_i, c_i)
            cell.text = text
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if r_i == 0 else WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.first_line_indent = None
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                for run in p.runs:
                    set_run_font(run, cn="宋体", en="Times New Roman", size_pt=10.5, bold=(r_i == 0))


def add_image_paragraph(doc: Document, img_path: Path):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    # 经验值：14cm 在 A4 正文宽度下较稳
    run.add_picture(str(img_path), width=Cm(14))
    set_exact_22(p)


def fill_from_markdown(template_path: Path, md_path: Path, output_path: Path):
    lines = md_path.read_text(encoding="utf-8").splitlines()
    data = extract_sections(lines)

    doc = Document(str(template_path))
    remove_body_from_anchor(doc, "论文/设计题目")
    enable_update_fields_on_open(doc)
    style_names = {s.name for s in doc.styles if s.type == WD_STYLE_TYPE.PARAGRAPH}

    # ---------- 标题 ----------
    if data["title_cn"]:
        p = doc.add_paragraph(data["title_cn"])
        style_title_cn(p)

    if data["title_en"]:
        p = doc.add_paragraph(data["title_en"])
        style_title_en(p)

    doc.add_paragraph()

    # ---------- 中文摘要 ----------
    p = doc.add_paragraph("摘  要")
    style_abstract_heading(p, "摘  要")

    for t in data["zh_abs_lines"]:
        p = doc.add_paragraph(t)
        style_body(p, indent=True)

    if data["zh_keywords"]:
        p = doc.add_paragraph(data["zh_keywords"])
        style_keywords(p)

    doc.add_page_break()

    # ---------- 英文摘要 ----------
    p = doc.add_paragraph("Abstract")
    style_abstract_heading(p, "Abstract")

    for t in data["en_abs_lines"]:
        p = doc.add_paragraph(t)
        style_body(p, indent=True)

    if data["en_keywords"]:
        p = doc.add_paragraph(data["en_keywords"])
        style_keywords(p)

    doc.add_page_break()

    # ---------- 目录 ----------
    p = doc.add_paragraph("目  录")
    # 使用模板已有目录标题样式
    if "目录 题目" in style_names:
        p.style = doc.styles["目录 题目"]
    elif "1-1级" in style_names:
        p.style = doc.styles["1-1级"]
    else:
        style_abstract_heading(p, "目  录")

    add_toc_field(doc)

    doc.add_page_break()

    # ---------- 正文 ----------
    body_lines = data["body_lines"]
    i = 0
    n = len(body_lines)
    base_dir = md_path.parent

    while i < n:
        raw = body_lines[i]
        line = raw.rstrip()
        stripped = line.strip()

        # 跳过空行和分隔线
        if not stripped or stripped == "---":
            i += 1
            continue

        # 跳过“使用网页模板生成截图”的操作说明
        if re.match(r"^[［\[]\s*使用\s+`[^`]+`\s+中图\d+.*[］\]]$", stripped):
            i += 1
            continue

        # 图片插入占位：［插入 `aGraduate/xxx.svg`］
        m_img = re.match(r"^[［\[]\s*插入\s+`([^`]+)`\s*[］\]]$", stripped)
        if m_img:
            img = resolve_image_path(m_img.group(1), base_dir)
            if img is not None and img.exists():
                add_image_paragraph(doc, img)
            i += 1
            continue

        # 标准 Markdown 图片：![alt](path)
        m_md_img = re.match(r"^!\[[^\]]*\]\(([^)]+)\)$", stripped)
        if m_md_img:
            img = resolve_image_path(m_md_img.group(1).strip(), base_dir)
            if img is not None and img.exists():
                add_image_paragraph(doc, img)
            i += 1
            continue

        # 表格
        if stripped.startswith("|") and (i + 1 < n) and is_table_separator(body_lines[i + 1].strip()):
            rows = [parse_markdown_table_row(body_lines[i])]
            i += 2  # 跳过表头和分隔线
            while i < n and body_lines[i].strip().startswith("|"):
                rows.append(parse_markdown_table_row(body_lines[i]))
                i += 1
            add_table(doc, rows)
            continue

        # 标题
        if stripped.startswith("# "):
            text = stripped[2:].strip()
            p = doc.add_paragraph(text)
            if "Heading 1" in style_names:
                p.style = doc.styles["Heading 1"]
            style_heading_chapter(p)
            i += 1
            continue
        if stripped.startswith("## "):
            text = stripped[3:].strip()
            p = doc.add_paragraph(text)
            if "Heading 2" in style_names:
                p.style = doc.styles["Heading 2"]
            style_heading_section(p)
            i += 1
            continue
        if stripped.startswith("### "):
            text = stripped[4:].strip()
            p = doc.add_paragraph(text)
            if "Heading 3" in style_names:
                p.style = doc.styles["Heading 3"]
            style_heading_subsection(p)
            i += 1
            continue
        if stripped.startswith("#### "):
            text = stripped[5:].strip()
            p = doc.add_paragraph(text)
            if "Heading 3" in style_names:
                p.style = doc.styles["Heading 3"]
            style_heading_subsection(p)
            i += 1
            continue

        # 图题 / 表题
        if re.match(r"^图\s*\d", stripped) or re.match(r"^图\d", stripped):
            p = doc.add_paragraph(stripped)
            style_caption(p, is_figure=True)
            i += 1
            continue
        if re.match(r"^表\s*\d", stripped) or re.match(r"^表\d", stripped):
            p = doc.add_paragraph(stripped)
            style_caption(p, is_figure=False)
            i += 1
            continue

        # 列表项
        if re.match(r"^[-*]\s+", stripped) or re.match(r"^\d+\.\s+", stripped):
            p = doc.add_paragraph(stripped)
            style_body(p, indent=False)
            i += 1
            continue

        # 参考文献条目
        if re.match(r"^\[\d+\]\s+", stripped):
            p = doc.add_paragraph(stripped)
            # 优先用模板中的参考文献样式
            if "参考文献 正文" in [s.name for s in doc.styles if s.type == WD_STYLE_TYPE.PARAGRAPH]:
                p.style = doc.styles["参考文献 正文"]
                for run in p.runs:
                    set_run_font(run, cn="宋体", en="Times New Roman", size_pt=12, bold=False)
            else:
                style_body(p, indent=False)
            i += 1
            continue

        # 普通正文
        p = doc.add_paragraph(stripped)
        style_body(p, indent=True)
        i += 1

    doc.save(str(output_path))


def main():
    base_dir = Path(__file__).resolve().parent
    md_path = base_dir / "06-毕业论文.md"
    template_path = base_dir / "2西南科技大学本科毕业论文（设计）模板2026版.docx"
    output_path = base_dir / "毕业论文_最终版.docx"

    if not md_path.exists():
        raise FileNotFoundError(f"找不到 Markdown 文件: {md_path}")
    if not template_path.exists():
        raise FileNotFoundError(f"找不到模板文件: {template_path}")

    fill_from_markdown(template_path, md_path, output_path)
    print(f"已生成文档: {output_path}")


if __name__ == "__main__":
    main()
