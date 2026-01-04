#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成研究思路文档 (docx格式)"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

def set_chinese_font(run, font_name='宋体', font_size=12):
    """设置中文字体"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def add_heading_cn(doc, text, level=1):
    """添加中文标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        set_chinese_font(run, '黑体', 16 if level == 1 else 14 if level == 2 else 12)
    return heading

def add_paragraph_cn(doc, text, bold=False, font_size=12):
    """添加中文段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_chinese_font(run, '宋体', font_size)
    run.bold = bold
    p.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进2字符
    p.paragraph_format.line_spacing = 1.5
    return p

def create_table(doc, headers, rows):
    """创建表格"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    # 表头
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                set_chinese_font(run, '宋体', 10)
    
    # 数据行
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row_data):
            row_cells[i].text = cell_text
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    set_chinese_font(run, '宋体', 10)
    
    return table

def main():
    doc = Document()
    
    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)
    
    # 标题
    title = doc.add_heading('拟采用的研究思路', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        set_chinese_font(run, '黑体', 18)
    
    # 副标题
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run('——研究方法、技术路线与可行性论证')
    set_chinese_font(subtitle_run, '黑体', 14)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # 空行
    
    # ==================== 一、研究方法 ====================
    add_heading_cn(doc, '一、研究方法', 1)
    
    add_heading_cn(doc, '1.1 总体研究方法', 2)
    add_paragraph_cn(doc, '本课题采用"理论研究—系统设计—原型实现—实验验证"的递进式研究方法，将工程实践与学术研究相结合。具体研究阶段如下表所示：')
    
    create_table(doc, 
        ['阶段', '方法', '产出'],
        [
            ['理论研究', '文献调研法、对比分析法', '技术选型报告、算法设计方案'],
            ['系统设计', '领域驱动设计(DDD)、微服务架构设计', '系统架构图、接口设计文档'],
            ['原型实现', '敏捷开发、迭代式开发', '可运行的系统原型'],
            ['实验验证', '对照实验法、压力测试法', '性能测试报告、算法评估报告'],
        ]
    )
    
    doc.add_paragraph()
    
    add_heading_cn(doc, '1.2 各研究目标的具体方法', 2)
    
    add_paragraph_cn(doc, '研究目标1：Golang云原生微服务架构', bold=True)
    add_paragraph_cn(doc, '采用架构模式研究与工程实践验证相结合的方法。首先调研业界主流微服务架构模式（Spring Cloud、Go-Micro、Kratos），对比分析RESTful与gRPC通信协议的性能差异；然后设计适配高并发场景的服务拆分策略；最后实现基于Gin框架的微服务集群。')
    
    add_paragraph_cn(doc, '研究目标2：多模态简历解析与逻辑风控', bold=True)
    add_paragraph_cn(doc, '采用算法融合研究与规则引擎设计相结合的方法。调研OCR技术（PaddleOCR、Tesseract）与NLP技术（NER、关系抽取），设计多模态文档解析流水线，构建基于规则的逻辑一致性校验引擎，通过标注数据集验证解析准确率。')
    
    add_paragraph_cn(doc, '研究目标3：可解释性人岗匹配算法', bold=True)
    add_paragraph_cn(doc, '采用深度学习与可解释AI研究相结合的方法。调研语义匹配模型（BERT、Sentence-BERT、双塔模型），设计多维度加权匹配算法，研究RAG技术在归因报告生成中的应用，通过A/B测试验证推荐效果。')
    
    add_paragraph_cn(doc, '研究目标4：高并发性能优化', bold=True)
    add_paragraph_cn(doc, '采用性能工程与压力测试验证相结合的方法。研究Goroutine调度模型与Channel通信机制，设计多级缓存策略（本地缓存+Redis），实现令牌桶限流与熔断降级机制，通过wrk/hey工具进行全链路压测。')

    # ==================== 二、技术路线 ====================
    add_heading_cn(doc, '二、技术路线', 1)
    
    add_heading_cn(doc, '2.1 总体技术路线', 2)
    add_paragraph_cn(doc, '本课题的技术路线分为三个阶段：')
    
    p = doc.add_paragraph()
    run = p.add_run('第一阶段（1-2月）：基础架构搭建。')
    set_chinese_font(run, '宋体', 12)
    run.bold = True
    run2 = p.add_run('完成微服务框架搭建、数据库设计、API网关实现、前端框架开发。')
    set_chinese_font(run2, '宋体', 12)
    
    p = doc.add_paragraph()
    run = p.add_run('第二阶段（3-4月）：核心算法实现。')
    set_chinese_font(run, '宋体', 12)
    run.bold = True
    run2 = p.add_run('完成简历解析引擎开发、匹配算法实现、RAG归因报告生成、风控规则引擎构建。')
    set_chinese_font(run2, '宋体', 12)
    
    p = doc.add_paragraph()
    run = p.add_run('第三阶段（5-6月）：性能优化验证。')
    set_chinese_font(run, '宋体', 12)
    run.bold = True
    run2 = p.add_run('完成压力测试、性能调优、指标验证、论文撰写。')
    set_chinese_font(run2, '宋体', 12)
    
    doc.add_paragraph()
    
    add_heading_cn(doc, '2.2 各模块技术路线', 2)
    
    add_paragraph_cn(doc, '（1）微服务架构技术路线', bold=True)
    add_paragraph_cn(doc, '需求分析→服务拆分→通信设计→服务治理→部署验证。采用Golang作为后端开发语言，Gin作为Web框架，实现8个业务微服务+1个API网关+1个AI评估服务的架构。当前已实现基于RESTful的服务通信和令牌桶限流，后续将引入gRPC提升服务间通信性能，引入Etcd实现服务注册发现。')
    
    add_paragraph_cn(doc, '（2）简历解析技术路线', bold=True)
    add_paragraph_cn(doc, '文件上传→格式识别→内容提取→结构化→风控校验→入库。当前已实现PDF/DOC/DOCX文件上传和基于正则的基础信息提取，后续将集成PaddleOCR实现图片简历识别，引入NER模型提升实体识别准确率，构建逻辑风控规则引擎。')
    
    add_paragraph_cn(doc, '（3）人岗匹配技术路线', bold=True)
    add_paragraph_cn(doc, '特征提取→向量化→相似度计算→多维加权→归因生成→推荐输出。当前已实现多维度匹配算法（技能50%+经验20%+位置15%+学历10%+薪资5%），后续将引入Sentence-BERT实现语义向量化，构建双塔模型提升语义匹配精度，集成RAG生成可解释的推荐归因报告。')
    
    add_paragraph_cn(doc, '（4）性能优化技术路线', bold=True)
    add_paragraph_cn(doc, '基准测试→瓶颈分析→优化实施→回归测试→指标验证。当前已实现Goroutine并发处理和令牌桶限流，后续将引入Redis多级缓存，实现Goroutine池化管理，进行数据库查询优化和全链路压测。')
    
    doc.add_paragraph()
    
    add_heading_cn(doc, '2.3 关键技术选型', 2)
    
    create_table(doc,
        ['技术领域', '选型方案', '选型理由'],
        [
            ['后端语言', 'Golang 1.21+', '原生并发支持、编译型高性能、云原生生态成熟'],
            ['Web框架', 'Gin', '轻量高性能、中间件丰富、社区活跃'],
            ['ORM框架', 'GORM', 'Go生态主流、功能完善、支持多数据库'],
            ['数据库', 'PostgreSQL', '支持数组类型、JSON、全文检索，适合复杂查询'],
            ['缓存', 'Redis', '高性能KV存储、支持多种数据结构、集群模式'],
            ['搜索引擎', 'Elasticsearch', '全文检索、日志分析、向量检索支持'],
            ['OCR引擎', 'PaddleOCR', '国产开源、中文识别优秀、部署灵活'],
            ['NLP模型', 'Sentence-BERT', '语义向量化、多语言支持、推理高效'],
            ['AI平台', 'Coze API', '工作流编排、快速集成、降低开发成本'],
            ['前端框架', 'Vue3 + TypeScript', '组合式API、类型安全、生态完善'],
            ['容器化', 'Docker + Compose', '标准化部署、环境一致性、易于扩展'],
        ]
    )

    # ==================== 三、可行性论证 ====================
    add_heading_cn(doc, '三、可行性论证', 1)
    
    add_heading_cn(doc, '3.1 技术可行性', 2)
    
    add_paragraph_cn(doc, '（1）微服务架构可行性', bold=True)
    add_paragraph_cn(doc, 'Golang微服务架构已在字节跳动、B站、七牛云等大厂广泛应用，技术栈成熟稳定。项目已实现8+1+1微服务架构，服务拆分合理，具备良好的扩展基础。Gin、GORM、Docker等工具链完善，开发效率高。风险评估为低风险，主要工作为性能优化和服务治理增强。')
    
    add_paragraph_cn(doc, '（2）简历解析可行性', bold=True)
    add_paragraph_cn(doc, 'PaddleOCR在中文场景准确率达95%+，支持PDF/图片多格式。命名实体识别(NER)技术成熟，开源模型（LAC、HanLP）可直接使用。项目已实现基于正则的基础解析和50+技能词库，可渐进式升级。风险评估为中等风险，需要模型训练和调优，但有开源方案兜底。')
    
    add_paragraph_cn(doc, '（3）人岗匹配算法可行性', bold=True)
    add_paragraph_cn(doc, '项目已实现多维度加权匹配算法，匹配逻辑清晰，可解释性强。Sentence-BERT等预训练模型可直接使用，无需从零训练。LangChain、LlamaIndex等RAG框架成熟，可快速集成。风险评估为中等风险，需要向量数据库和模型推理资源。')
    
    add_paragraph_cn(doc, '（4）高并发性能可行性', bold=True)
    add_paragraph_cn(doc, 'Golang原生支持Goroutine，单机可轻松支撑百万级并发。项目已实现令牌桶限流、数据库连接池，具备优化基础。4核8G环境下，通过缓存和异步优化可达成1000 QPS。风险评估为低风险，Golang性能优化方案成熟，有大量最佳实践。')
    
    add_heading_cn(doc, '3.2 经济可行性', 2)
    
    create_table(doc,
        ['资源类型', '需求', '成本估算', '说明'],
        [
            ['开发环境', 'MacBook Pro', '已有', '本地开发测试'],
            ['服务器', '4核8G云服务器', '~100元/月', '阿里云/腾讯云学生优惠'],
            ['数据库', 'PostgreSQL', '免费', '开源数据库'],
            ['AI服务', 'Coze API', '免费额度', '字节跳动免费额度充足'],
            ['OCR服务', 'PaddleOCR', '免费', '百度开源，本地部署'],
            ['向量模型', 'Sentence-BERT', '免费', 'HuggingFace开源模型'],
        ]
    )
    
    add_paragraph_cn(doc, '总成本估算约100-200元/月，在学生可承受范围内。')
    
    add_heading_cn(doc, '3.3 时间可行性', 2)
    
    create_table(doc,
        ['阶段', '时间', '主要任务', '产出物'],
        [
            ['第一阶段', '第1-2月', '基础架构完善、文献调研', '架构优化、开题报告'],
            ['第二阶段', '第3-4月', '核心算法实现', '简历解析引擎、匹配算法'],
            ['第三阶段', '第5月', '性能优化与测试', '压测报告、性能指标'],
            ['第四阶段', '第6月', '论文撰写与答辩准备', '毕业论文、演示系统'],
        ]
    )
    
    add_paragraph_cn(doc, '当前进度评估：微服务架构已搭建完成（8+1+1服务），前端界面已开发完成（Vue3 + Element Plus），基础匹配算法已实现，简历上传解析基础功能已实现。项目已完成约60%的基础工作，剩余工作量在4-5个月内可完成。')
    
    add_heading_cn(doc, '3.4 风险评估与应对', 2)
    
    create_table(doc,
        ['风险类型', '风险描述', '概率', '影响', '应对措施'],
        [
            ['技术风险', 'OCR识别准确率不达标', '中', '高', '使用PaddleOCR预训练模型，必要时人工标注微调'],
            ['技术风险', '语义匹配模型推理慢', '中', '中', '模型量化、批量推理、结果缓存'],
            ['资源风险', '服务器资源不足', '低', '中', '优化算法复杂度、使用轻量级模型'],
            ['时间风险', '开发进度延迟', '中', '高', '优先实现核心功能，非核心功能降级处理'],
            ['数据风险', '测试数据不足', '低', '中', '使用公开简历数据集+模拟数据生成'],
        ]
    )

    # ==================== 四、预期成果 ====================
    add_heading_cn(doc, '四、预期成果', 1)
    
    add_heading_cn(doc, '4.1 系统成果', 2)
    add_paragraph_cn(doc, '（1）可运行的微服务人才运营平台：包含8个业务微服务+API网关+AI评估服务，完整的前端管理界面，Docker一键部署方案。')
    add_paragraph_cn(doc, '（2）核心算法模块：多模态简历解析引擎（OCR+NLP）、可解释性人岗匹配算法（语义向量+RAG归因）、逻辑风控规则引擎。')
    
    add_heading_cn(doc, '4.2 量化指标', 2)
    
    create_table(doc,
        ['指标', '目标值', '验证方法'],
        [
            ['简历解析准确率', '≥ 95%', '标注数据集测试'],
            ['推荐准确率 Precision@10', '≥ 80%', 'A/B测试 + 人工评估'],
            ['系统吞吐量', '≥ 1000 QPS', 'wrk/hey 压力测试'],
            ['平均响应延迟', '< 300ms', '全链路监控'],
            ['P99 响应延迟', '< 500ms', '压测统计'],
        ]
    )
    
    add_heading_cn(doc, '4.3 学术成果', 2)
    add_paragraph_cn(doc, '（1）毕业论文：《基于Golang微服务架构的智能人才运营平台设计与实现》')
    add_paragraph_cn(doc, '（2）技术文档：系统架构文档、API接口文档、部署文档')
    add_paragraph_cn(doc, '（3）开源代码：完整的项目源码（可选）')

    # ==================== 五、总结 ====================
    add_heading_cn(doc, '五、总结', 1)
    
    add_paragraph_cn(doc, '本课题基于已有的微服务人才运营平台，通过引入OCR+NLP融合的简历解析技术、基于语义向量的可解释性匹配算法、以及Golang高并发优化策略，实现对现有系统的全面升级。')
    
    add_paragraph_cn(doc, '核心创新点：（1）突破"匹配黑盒"，实现推荐结果的白盒化与可解释化；（2）融合OCR与NLP技术，解决多模态简历的精准量化问题；（3）验证Golang微服务在有限算力下的工业级性能。')
    
    add_paragraph_cn(doc, '可行性结论：技术可行——核心技术栈成熟，开源方案丰富；经济可行——成本可控，学生可承受；时间可行——基础工作已完成60%，剩余工作量合理。')
    
    # 保存文档
    doc.save('docs/研究思路_方法技术路线可行性论证.docx')
    print('文档已生成: docs/研究思路_方法技术路线可行性论证.docx')

if __name__ == '__main__':
    main()
