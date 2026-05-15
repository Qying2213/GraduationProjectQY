#!/usr/bin/env python3
"""
Batch-fix common WPS/Word thesis format issues for the local .docx file.

The script creates a corrected copy by default. It does not overwrite the
original document unless --in-place is passed explicitly.
"""

from __future__ import annotations

import argparse
import re
import shutil
import tempfile
from pathlib import Path
from zipfile import ZipFile, ZIP_DEFLATED

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


BASE_DIR = Path(__file__).resolve().parent
DEFAULT_INPUT = BASE_DIR / "微服务架构的智能人才运营平台设计与实现-秦洋-计科2206-5120227635 (1).docx"
DEFAULT_OUTPUT_SUFFIX = "_格式修正版"

FONT_CN_BODY = "宋体"
FONT_CN_HEADING = "黑体"
FONT_EN = "Times New Roman"

SIZE_XIAOYI = Pt(24)
SIZE_XIAOER = Pt(18)
SIZE_XIAOSAN = Pt(15)
SIZE_SIHAO = Pt(14)
SIZE_XIAOSI = Pt(12)
SIZE_WUHAO = Pt(10.5)
LINE_22 = Pt(22)
LINE_30 = Pt(30)

CHAPTER_RE = re.compile(r"^[1-6](?:\s+[\u4e00-\u9fff]|[\u4e00-\u9fff])")
SECTION_RE = re.compile(r"^[1-9]\d*(?:\.[1-9]\d*){1,3}\s*\S+")
REFERENCE_RE = re.compile(r"^\[\d+\]")
CAPTION_RE = re.compile(r"^[\x00-\x1f\s]*(图|表)\s*([1-9]\d*)-(?:(0)-)?([1-9]\d*)([A-Za-z]?)\s*(.*)?$")


class Stats(dict):
    def add(self, key: str, value: int = 1) -> None:
        self[key] = self.get(key, 0) + value


def get_or_add_child(parent, tag: str):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def remove_attr(element, attr: str) -> None:
    qattr = qn(attr)
    if qattr in element.attrib:
        del element.attrib[qattr]


def set_style_font(style, cn_font: str, size=None, bold=None) -> None:
    if style is None:
        return
    font = style.font
    font.name = FONT_EN
    if size is not None:
        font.size = size
    if bold is not None:
        font.bold = bold

    rpr = style.element.get_or_add_rPr()
    rfonts = get_or_add_child(rpr, "w:rFonts")
    for key, value in (
        ("w:ascii", FONT_EN),
        ("w:hAnsi", FONT_EN),
        ("w:cs", FONT_EN),
        ("w:eastAsia", cn_font),
    ):
        rfonts.set(qn(key), value)


def set_run_font(run, cn_font: str, size=None, bold=None, italic=False) -> None:
    font = run.font
    font.name = FONT_EN
    if size is not None:
        font.size = size
    if bold is not None:
        font.bold = bold
    if italic is not None:
        font.italic = italic

    rpr = run._element.get_or_add_rPr()
    rfonts = get_or_add_child(rpr, "w:rFonts")
    for key, value in (
        ("w:ascii", FONT_EN),
        ("w:hAnsi", FONT_EN),
        ("w:cs", FONT_EN),
        ("w:eastAsia", cn_font),
    ):
        rfonts.set(qn(key), value)


def set_line_spacing_exact_22(paragraph) -> None:
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = LINE_22


def set_line_spacing_exact(paragraph, value) -> None:
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = value


def paragraph_has_drawing(paragraph) -> bool:
    xml = paragraph._element.xml
    return "<w:drawing" in xml or "<w:pict" in xml


def set_spacing_lines(paragraph, before_lines: int | None = None, after_lines: int | None = None) -> None:
    """Set paragraph spacing with Word's line-unit attributes.

    WPS thesis checkers often distinguish "0.5 line" from equivalent point
    values, so write w:beforeLines/w:afterLines directly.
    """
    ppr = paragraph._p.get_or_add_pPr()
    spacing = get_or_add_child(ppr, "w:spacing")

    for attr in ("w:before", "w:after"):
        remove_attr(spacing, attr)

    if before_lines is None:
        remove_attr(spacing, "w:beforeLines")
    else:
        spacing.set(qn("w:beforeLines"), str(before_lines))

    if after_lines is None:
        remove_attr(spacing, "w:afterLines")
    else:
        spacing.set(qn("w:afterLines"), str(after_lines))


def set_spacing_points(paragraph, before_pt: float = 0, after_pt: float = 0) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    spacing = get_or_add_child(ppr, "w:spacing")

    for attr in ("w:beforeLines", "w:afterLines"):
        remove_attr(spacing, attr)

    spacing.set(qn("w:before"), str(round(before_pt * 20)))
    spacing.set(qn("w:after"), str(round(after_pt * 20)))


def set_first_line_indent_chars(paragraph, chars: int | None) -> None:
    """Use character-unit first-line indent, e.g. 200 means 2.00 chars."""
    ppr = paragraph._p.get_or_add_pPr()
    ind = get_or_add_child(ppr, "w:ind")

    for attr in ("w:firstLine", "w:hanging", "w:hangingChars"):
        remove_attr(ind, attr)

    if chars is None:
        remove_attr(ind, "w:firstLineChars")
    else:
        ind.set(qn("w:firstLineChars"), str(chars))


def set_paragraph_base(
    paragraph,
    *,
    alignment=None,
    first_line_chars: int | None = None,
    spacing_lines: tuple[int | None, int | None] | None = None,
    spacing_points: tuple[float, float] | None = None,
) -> None:
    if alignment is not None:
        paragraph.alignment = alignment
    # Do not force exact 22pt line spacing on paragraphs that contain inline
    # pictures. WPS clips large inline pictures to the fixed line height, which
    # looks like the picture disappeared and text overlaps the image area.
    if not paragraph_has_drawing(paragraph):
        set_line_spacing_exact_22(paragraph)
    set_first_line_indent_chars(paragraph, first_line_chars)

    if spacing_lines is not None:
        set_spacing_lines(paragraph, spacing_lines[0], spacing_lines[1])
    elif spacing_points is not None:
        set_spacing_points(paragraph, spacing_points[0], spacing_points[1])


def apply_font_to_paragraph(paragraph, cn_font: str, size, bold=None) -> None:
    for run in paragraph.runs:
        set_run_font(run, cn_font, size=size, bold=bold)


def has_page_break(paragraph) -> bool:
    return 'w:type="page"' in paragraph._element.xml


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def set_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def make_page_break_paragraph():
    paragraph = OxmlElement("w:p")
    run = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run.append(br)
    paragraph.append(run)
    return paragraph


def make_toc_field_paragraph():
    paragraph = OxmlElement("w:p")

    def append_run(child):
        run = OxmlElement("w:r")
        run.append(child)
        paragraph.append(run)

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    append_run(begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r' TOC \o "1-3" \h \z \u '
    append_run(instr)

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    append_run(separate)

    text = OxmlElement("w:t")
    text.text = "请在 WPS 中右键更新目录"
    append_run(text)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    append_run(end)

    return paragraph


def insert_toc_field_if_missing(doc: Document, stats: Stats) -> None:
    """Insert a TOC field when the document has only a TOC title.

    The field still needs to be updated in WPS/Word, but it prevents the main
    chapter text from sitting directly under the TOC title.
    """
    toc_idx = find_paragraph_index(doc, lambda text: re.sub(r"\s+", "", text) == "目录")
    if toc_idx is None:
        return

    paragraphs = doc.paragraphs
    first_after_toc_idx = None
    for idx in range(toc_idx + 1, len(paragraphs)):
        if paragraphs[idx].text.strip():
            first_after_toc_idx = idx
            break

    if first_after_toc_idx is None:
        return

    first_after_toc = paragraphs[first_after_toc_idx]
    if "TOC" in (first_after_toc.style.name if first_after_toc.style else ""):
        return
    if "\t" in first_after_toc.text:
        return
    if not (first_after_toc.style and first_after_toc.style.name.startswith("Heading")):
        return

    body = first_after_toc._element.getparent()
    insert_at = body.index(first_after_toc._element)
    body.insert(insert_at, make_toc_field_paragraph())
    body.insert(insert_at + 1, make_page_break_paragraph())
    stats.add("inserted_toc_field")


def find_paragraph_index(doc: Document, predicate) -> int | None:
    for idx, paragraph in enumerate(doc.paragraphs):
        if predicate(paragraph.text.strip()):
            return idx
    return None


def remove_blank_paragraphs_between(doc: Document, start_idx: int | None, end_idx: int | None, stats: Stats) -> None:
    if start_idx is None or end_idx is None or end_idx <= start_idx:
        return

    for paragraph in list(doc.paragraphs[start_idx + 1 : end_idx]):
        if not paragraph.text.strip() and not has_page_break(paragraph) and not paragraph_has_drawing(paragraph):
            remove_paragraph(paragraph)
            stats.add("removed_blank_abstract_paragraphs")


def ensure_blank_before(paragraph, stats: Stats) -> None:
    previous = paragraph._element.getprevious()
    if previous is not None and previous.tag.endswith("}p"):
        text_nodes = previous.xpath(".//w:t/text()")
        has_text = bool("".join(text_nodes).strip())
        if not has_text:
            return

    new_paragraph = paragraph.insert_paragraph_before("")
    set_paragraph_base(
        new_paragraph,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_chars=None,
        spacing_points=(0, 0),
    )
    stats.add("inserted_blank_before_english_keywords")


def normalize_styles(doc: Document) -> None:
    styles = doc.styles
    for name in ("Normal", "Normal (Web)", "Body Text", "Body Text Indent 2"):
        if name in styles:
            set_style_font(styles[name], FONT_CN_BODY, SIZE_XIAOSI, bold=False)
    for name in ("TOC 1", "TOC 2", "TOC 3", "目录 1", "目录 2", "目录 3"):
        if name in styles:
            set_style_font(styles[name], FONT_CN_BODY, SIZE_XIAOSI, bold=False)
    if "参考文献 正文" in styles:
        set_style_font(styles["参考文献 正文"], FONT_CN_BODY, SIZE_WUHAO, bold=False)
    for name in ("Heading 1", "Heading 2", "Heading 3"):
        if name in styles:
            set_style_font(styles[name], FONT_CN_HEADING, None, bold=False)


def format_abstract_title(paragraph, *, english: bool, stats: Stats) -> None:
    if not english:
        set_paragraph_text(paragraph, "摘要：")
    paragraph.style = paragraph.style
    set_paragraph_base(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_chars=None,
        spacing_lines=(50, 50),
    )
    apply_font_to_paragraph(
        paragraph,
        FONT_EN if english else FONT_CN_HEADING,
        SIZE_XIAOYI,
        bold=False,
    )
    stats.add("formatted_abstract_titles")


def format_keywords(paragraph, *, english: bool, stats: Stats) -> None:
    set_paragraph_base(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.LEFT if english else WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_chars=None,
        spacing_points=(0, 0) if english else None,
        spacing_lines=None if english else (50, 50),
    )
    apply_font_to_paragraph(
        paragraph,
        FONT_EN if english else FONT_CN_HEADING,
        SIZE_XIAOSI,
        bold=False,
    )
    stats.add("formatted_keywords")


def format_chapter_heading(paragraph, stats: Stats) -> None:
    set_paragraph_base(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_chars=None,
        spacing_lines=(50, 50),
    )
    apply_font_to_paragraph(paragraph, FONT_CN_HEADING, SIZE_XIAOER, bold=False)
    stats.add("formatted_chapter_headings")


def format_section_heading(paragraph, level: int, stats: Stats) -> None:
    size = SIZE_SIHAO if level == 2 else SIZE_XIAOSI
    set_paragraph_base(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_chars=None,
        spacing_lines=(50, 50),
    )
    apply_font_to_paragraph(paragraph, FONT_CN_HEADING, size, bold=False)
    stats.add(f"formatted_level_{level}_headings")


def format_reference_title(paragraph, stats: Stats) -> None:
    set_paragraph_base(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_chars=None,
        spacing_lines=(50, 50),
    )
    apply_font_to_paragraph(paragraph, FONT_CN_HEADING, SIZE_XIAOER, bold=False)
    stats.add("formatted_reference_title")


def format_reference_item(paragraph, stats: Stats) -> None:
    set_paragraph_base(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_chars=None,
        spacing_points=(0, 0),
    )
    apply_font_to_paragraph(paragraph, FONT_CN_BODY, SIZE_WUHAO, bold=False)
    stats.add("formatted_reference_items")


def format_body(paragraph, stats: Stats) -> None:
    set_paragraph_base(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_chars=200,
        spacing_points=(0, 0),
    )
    apply_font_to_paragraph(paragraph, FONT_CN_BODY, SIZE_XIAOSI, bold=False)
    stats.add("formatted_body_paragraphs")


def format_chinese_abstract_body(paragraph, stats: Stats) -> None:
    set_paragraph_base(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_chars=None,
        spacing_points=(0, 0),
    )
    apply_font_to_paragraph(paragraph, FONT_CN_BODY, SIZE_XIAOSI, bold=False)
    stats.add("formatted_chinese_abstract_body")


def format_english_abstract_body(paragraph, stats: Stats) -> None:
    set_paragraph_base(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_chars=None,
        spacing_points=(0, 0),
    )
    apply_font_to_paragraph(paragraph, FONT_EN, SIZE_XIAOSI, bold=False)
    stats.add("formatted_english_abstract_body")


def format_toc_title(paragraph, stats: Stats) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    set_first_line_indent_chars(paragraph, None)
    set_spacing_points(paragraph, 0, 0)
    apply_font_to_paragraph(paragraph, FONT_CN_HEADING, SIZE_XIAOER, bold=False)
    stats.add("formatted_toc_title")


def format_toc_entry(paragraph, stats: Stats) -> None:
    set_paragraph_base(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_chars=None,
        spacing_points=(0, 0),
    )
    apply_font_to_paragraph(paragraph, FONT_CN_BODY, SIZE_XIAOSI, bold=False)
    stats.add("formatted_toc_entries")


def normalize_caption_text(paragraph, stats: Stats) -> str:
    text = paragraph.text.strip().lstrip("\x01")
    match = CAPTION_RE.match(text)
    if not match:
        return text

    kind, chapter, zero_part, number, letter, title = match.groups()
    # The checker expects figure/table numbers like 图4-1 rather than 图4-0-1.
    if zero_part:
        stats.add("normalized_zero_caption_numbers")
    title = (title or "").strip()
    # Letter suffixes such as 图5-3a are kept as content only when necessary;
    # the checker generally expects the core number and title to be separated.
    if letter and title:
        title = f"{letter}{title}"
    normalized = f"{kind}{chapter}-{number}"
    if title:
        normalized += f" {title}"
    if normalized != text:
        set_paragraph_text(paragraph, normalized)
        stats.add("normalized_caption_text")
    return normalized


def format_caption(paragraph, stats: Stats) -> None:
    text = normalize_caption_text(paragraph, stats)
    kind = text[:1]
    set_paragraph_base(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_chars=None,
        spacing_points=(0, 0),
    )
    apply_font_to_paragraph(
        paragraph,
        FONT_CN_BODY,
        SIZE_WUHAO,
        bold=True if kind == "表" else False,
    )
    stats.add("formatted_captions")


def collect_and_renumber_captions(doc: Document, stats: Stats) -> dict[str, str]:
    counters: dict[tuple[str, str], int] = {}
    replacements: dict[str, str] = {}

    for paragraph in doc.paragraphs:
        raw_text = paragraph.text.strip().lstrip("\x01")
        match = CAPTION_RE.match(raw_text)
        if not match:
            continue

        kind, chapter, zero_part, number, letter, title = match.groups()
        title = (title or "").strip()
        key = (kind, chapter)
        counters[key] = counters.get(key, 0) + 1
        new_label = f"{kind}{chapter}-{counters[key]}"
        old_label = f"{kind}{chapter}-{'0-' if zero_part else ''}{number}{letter}"
        normalized = new_label if not title else f"{new_label} {title}"

        if normalized != raw_text:
            set_paragraph_text(paragraph, normalized)
            stats.add("renumbered_captions")

        for variant in {
            old_label,
            old_label.replace(kind, f"{kind} ", 1),
            old_label.replace("-", "－"),
        }:
            replacements[variant] = new_label

    return {old: new for old, new in replacements.items() if old != new}


def replace_caption_references(doc: Document, replacements: dict[str, str], stats: Stats) -> None:
    if not replacements:
        return

    for paragraph in doc.paragraphs:
        if CAPTION_RE.match(paragraph.text.strip().lstrip("\x01")):
            continue
        for run in paragraph.runs:
            original = run.text
            updated = original
            for old, new in replacements.items():
                updated = updated.replace(old, new)
            if updated != original:
                run.text = updated
                stats.add("updated_caption_references")


def format_table_paragraph(paragraph, stats: Stats) -> None:
    if paragraph_has_drawing(paragraph):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        stats.add("preserved_table_picture_paragraphs")
        return
    if not paragraph.text.strip():
        return
    set_paragraph_base(
        paragraph,
        alignment=paragraph.alignment,
        first_line_chars=None,
        spacing_points=(0, 0),
    )
    apply_font_to_paragraph(paragraph, FONT_CN_BODY, SIZE_XIAOSI, bold=False)
    stats.add("formatted_table_paragraphs")


def normalize_reference_punctuation(paragraph, stats: Stats) -> None:
    """Conservative punctuation fix for references ending with Chinese full stop."""
    text = paragraph.text.strip()
    if not REFERENCE_RE.match(text) or not text.endswith("。"):
        return

    # Modify only the last textual run, keeping the rest of the reference intact.
    for run in reversed(paragraph.runs):
        if run.text:
            stripped = run.text.rstrip()
            if stripped.endswith("。"):
                run.text = run.text[: len(run.text) - len(stripped)] + stripped[:-1] + "."
                stats.add("normalized_reference_periods")
            return


def patch_cover_textboxes(docx_path: Path) -> None:
    """Patch cover fields stored in text boxes/shapes.

    python-docx does not expose every WPS text box as a normal paragraph, while
    the format checker still validates those cover fields. Patch the underlying
    WordprocessingML for the exact cover labels and values reported by checker.
    """
    cover_texts = {
        "学院名称",
        "计算机科学与技术学院",
        "专业名称",
        "计算机科学与技术",
        "学生姓名",
        "秦洋",
        "学号",
        "5120227635",
        "指导教师",
        "李胤贤",
    }
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    w = f"{{{ns}}}"

    with tempfile.TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)
        with ZipFile(docx_path, "r") as source:
            source.extractall(tmp_path)

        document_xml = tmp_path / "word" / "document.xml"
        try:
            from lxml import etree
        except ImportError:
            return

        root = etree.fromstring(document_xml.read_bytes())
        namespaces = {"w": ns}

        for paragraph in root.xpath(".//w:p", namespaces=namespaces):
            text = "".join(paragraph.xpath(".//w:t/text()", namespaces=namespaces)).strip()
            if text not in cover_texts:
                continue

            ppr = paragraph.find(f"{w}pPr")
            if ppr is None:
                ppr = etree.Element(f"{w}pPr")
                paragraph.insert(0, ppr)

            spacing = ppr.find(f"{w}spacing")
            if spacing is None:
                spacing = etree.SubElement(ppr, f"{w}spacing")
            for attr in ("line", "lineRule", "before", "after", "beforeLines", "afterLines"):
                spacing.attrib.pop(f"{w}{attr}", None)
            spacing.set(f"{w}line", "600")
            spacing.set(f"{w}lineRule", "exact")

            for run in paragraph.xpath(".//w:r", namespaces=namespaces):
                rpr = run.find(f"{w}rPr")
                if rpr is None:
                    rpr = etree.Element(f"{w}rPr")
                    run.insert(0, rpr)

                rfonts = rpr.find(f"{w}rFonts")
                if rfonts is None:
                    rfonts = etree.SubElement(rpr, f"{w}rFonts")
                rfonts.set(f"{w}ascii", FONT_EN)
                rfonts.set(f"{w}hAnsi", FONT_EN)
                rfonts.set(f"{w}cs", FONT_EN)
                rfonts.set(f"{w}eastAsia", FONT_CN_HEADING)

                for tag, value in (("sz", "30"), ("szCs", "30")):
                    elem = rpr.find(f"{w}{tag}")
                    if elem is None:
                        elem = etree.SubElement(rpr, f"{w}{tag}")
                    elem.set(f"{w}val", value)

        document_xml.write_bytes(
            etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
        )

        temp_docx = docx_path.with_suffix(docx_path.suffix + ".tmp")
        with ZipFile(temp_docx, "w", ZIP_DEFLATED) as target:
            for file_path in tmp_path.rglob("*"):
                if file_path.is_file():
                    target.write(file_path, file_path.relative_to(tmp_path).as_posix())
        temp_docx.replace(docx_path)


def fix_document(doc: Document, *, fix_reference_periods: bool = False) -> Stats:
    stats = Stats()
    normalize_styles(doc)
    insert_toc_field_if_missing(doc, stats)
    caption_replacements = collect_and_renumber_captions(doc, stats)
    replace_caption_references(doc, caption_replacements, stats)

    abstract_idx = find_paragraph_index(doc, lambda text: text.replace(" ", "") in {"摘要", "摘要："})
    cn_keywords_idx = find_paragraph_index(doc, lambda text: text.startswith("关键词"))
    english_abstract_idx = find_paragraph_index(doc, lambda text: text.lower() == "abstract")
    english_keywords_idx = find_paragraph_index(doc, lambda text: text.lower().startswith(("key words", "keywords")))

    remove_blank_paragraphs_between(doc, abstract_idx, cn_keywords_idx, stats)

    abstract_idx = find_paragraph_index(doc, lambda text: text.replace(" ", "") in {"摘要", "摘要："})
    cn_keywords_idx = find_paragraph_index(doc, lambda text: text.startswith("关键词"))
    english_abstract_idx = find_paragraph_index(doc, lambda text: text.lower() == "abstract")
    english_keywords_idx = find_paragraph_index(doc, lambda text: text.lower().startswith(("key words", "keywords")))
    if english_keywords_idx is not None:
        ensure_blank_before(doc.paragraphs[english_keywords_idx], stats)
    english_keywords_idx = find_paragraph_index(doc, lambda text: text.lower().startswith(("key words", "keywords")))

    seen_reference_title = False
    in_references = False
    in_toc = False

    for para_idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        compact = re.sub(r"\s+", "", text)
        if paragraph_has_drawing(paragraph):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            stats.add("preserved_picture_paragraphs")
            continue
        if not text:
            continue

        if compact in {"摘要", "摘要："}:
            format_abstract_title(paragraph, english=False, stats=stats)
        elif text.lower() == "abstract":
            format_abstract_title(paragraph, english=True, stats=stats)
        elif text.startswith("关键词"):
            format_keywords(paragraph, english=False, stats=stats)
        elif text.lower().startswith(("key words", "keywords")):
            format_keywords(paragraph, english=True, stats=stats)
        elif (
            abstract_idx is not None
            and cn_keywords_idx is not None
            and abstract_idx < para_idx < cn_keywords_idx
        ):
            format_chinese_abstract_body(paragraph, stats)
        elif (
            english_abstract_idx is not None
            and english_keywords_idx is not None
            and english_abstract_idx < para_idx < english_keywords_idx
        ):
            format_english_abstract_body(paragraph, stats)
        elif compact == "目录":
            format_toc_title(paragraph, stats)
            in_toc = True
        elif in_toc and paragraph.style and paragraph.style.name.startswith("Heading"):
            in_toc = False
            if CHAPTER_RE.match(text) and "." not in text.split()[0]:
                format_chapter_heading(paragraph, stats)
            elif SECTION_RE.match(text):
                level = min(text.split()[0].count(".") + 1, 4)
                format_section_heading(paragraph, level, stats)
        elif in_toc or (paragraph.style and paragraph.style.name.startswith(("TOC", "目录"))) or "\t" in text:
            format_toc_entry(paragraph, stats)
        elif CAPTION_RE.match(text):
            format_caption(paragraph, stats)
        elif text == "参考文献":
            format_reference_title(paragraph, stats)
            seen_reference_title = True
            in_references = True
        elif in_references and REFERENCE_RE.match(text):
            format_reference_item(paragraph, stats)
            if fix_reference_periods:
                normalize_reference_punctuation(paragraph, stats)
        elif CHAPTER_RE.match(text) and "." not in text.split()[0]:
            format_chapter_heading(paragraph, stats)
            in_references = False
        elif SECTION_RE.match(text):
            level = min(text.split()[0].count(".") + 1, 4)
            format_section_heading(paragraph, level, stats)
            in_references = False
        elif paragraph.style and paragraph.style.name in {"Normal", "Normal (Web)", "Body Text", "Body Text Indent 2"}:
            # Skip early cover/authorization text before the abstract title. It uses
            # many manual positions and should be adjusted in the official template.
            if abstract_idx is None or para_idx > abstract_idx:
                format_body(paragraph, stats)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    format_table_paragraph(paragraph, stats)

    if not seen_reference_title:
        stats.add("warning_reference_title_not_found", 0)

    return stats


def build_output_path(input_path: Path, output_arg: str | None, in_place: bool) -> Path:
    if in_place:
        return input_path
    if output_arg:
        return Path(output_arg).expanduser().resolve()
    return input_path.with_name(f"{input_path.stem}{DEFAULT_OUTPUT_SUFFIX}{input_path.suffix}")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="批量修正西南科技大学本科论文 docx 常见格式问题。")
    parser.add_argument(
        "input",
        nargs="?",
        default=str(DEFAULT_INPUT),
        help="要处理的 .docx 文件，默认处理本目录下的论文文件。",
    )
    parser.add_argument("-o", "--output", help="输出文件路径。默认生成 *_格式修正版.docx。")
    parser.add_argument("--in-place", action="store_true", help="直接覆盖原文件；会先生成 .bak 备份。")
    parser.add_argument(
        "--fix-reference-periods",
        action="store_true",
        help="把参考文献末尾的中文句号保守替换为英文句点。",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    input_path = Path(args.input).expanduser().resolve()
    if not input_path.exists():
        raise FileNotFoundError(f"找不到输入文件：{input_path}")

    output_path = build_output_path(input_path, args.output, args.in_place)
    if args.in_place:
        backup_path = input_path.with_suffix(input_path.suffix + ".bak")
        shutil.copy2(input_path, backup_path)
        print(f"已备份原文件：{backup_path}")

    doc = Document(input_path)
    stats = fix_document(doc, fix_reference_periods=args.fix_reference_periods)
    doc.save(output_path)
    patch_cover_textboxes(output_path)

    print(f"已生成：{output_path}")
    print("处理统计：")
    for key in sorted(stats):
        print(f"  {key}: {stats[key]}")
    print("\n建议：用 WPS 打开修正版后，右键目录选择“更新域/更新目录”，再重新跑格式检测。")


if __name__ == "__main__":
    main()
